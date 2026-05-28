import fitz  # PyMuPDF
import io
import re
import docx

def clean_pdf_text(text):
    """
    Remove URLs, and normalize whitespace in extracted PDF text.
    """
    if not text:
        return ""
    # Remove URLs
    text = re.sub(r'https?://\S+|www\.\S+', ' ', text)
    # Replace newlines and carriage returns with spaces
    text = text.replace('\n', ' ').replace('\r', ' ')
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_text_from_pdf_bytes(pdf_bytes):
    """
    Extracts plain text from PDF data provided as bytes.
    """
    if not pdf_bytes:
        return ""
    
    try:
        # Open PDF from bytes
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return clean_pdf_text(text)
    except Exception as e:
        print(f"Error extracting text from PDF bytes: {e}")
        return ""

def extract_text_from_pdf_file(file_path):
    """
    Extracts plain text from a PDF file on disk.
    """
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return clean_pdf_text(text)
    except Exception as e:
        print(f"Error extracting text from PDF file: {e}")
        return ""

def extract_text_from_docx_bytes(docx_bytes):
    """
    Extracts plain text from a DOCX document provided as bytes.
    """
    if not docx_bytes:
        return ""
    
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return clean_pdf_text('\n'.join(full_text))
    except Exception as e:
        print(f"Error extracting text from DOCX bytes: {e}")
        return ""

def extract_text_from_docx_file(file_path):
    """
    Extracts plain text from a DOCX file on disk.
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return clean_pdf_text('\n'.join(full_text))
    except Exception as e:
        print(f"Error extracting text from DOCX file: {e}")
        return ""
