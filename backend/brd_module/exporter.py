"""
exporter.py
Compiles the latest version of all BRD sections into a single Markdown document,
including a section for Validation Flags. Supports Markdown, PDF, and DOCX export
with template-based formatting for DOCX.
"""
from brd_module.storage import get_latest_brd_sections, get_connection
from datetime import datetime, timezone
import markdown
import re
import os
from io import BytesIO

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    # OSError occurs on Windows when native GLib/Pango DLLs are missing.
    # PDF export will be unavailable but all other endpoints still work.
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

def export_brd(session_id: str, title: str = "Business Requirements Document") -> str:
    """
    Fetches the latest BRD sections and any active validation flags,
    and returns a formatted Markdown string.
    """
    sections = get_latest_brd_sections(session_id)
    
    # Order of presentation
    section_order = [
        ("executive_summary", "1. Executive Summary"),
        ("functional_requirements", "2. Functional Requirements"),
        ("stakeholder_analysis", "3. Stakeholder Analysis"),
        ("timeline", "4. Project Timeline"),
        ("decisions", "5. Key Decisions"),
        ("assumptions", "6. Assertions & Assumptions"),
        ("success_metrics", "7. Success Metrics")
    ]
    
    doc = []
    doc.append(f"# {title}")
    doc.append(f"**Generated:** {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
    doc.append(f"**Session ID:** `{session_id}`")
    doc.append("---\n")
    
    # 1. Fetch Validation Flags
    conn, db_type = get_connection()
    flags = []
    try:
        cur = conn.cursor()
        query = """
            SELECT section_name, flag_type, severity, description 
            FROM brd_validation_flags 
            WHERE session_id = %s
            ORDER BY severity DESC
        """
        if db_type == "sqlite": query = query.replace("%s", "?")
        
        cur.execute(query, (session_id,))
        flags = cur.fetchall()
    except Exception as e:
        doc.append(f"> **Warning:** Could not fetch validation flags: {e}\n")
    finally:
        conn.close()
        
    if flags:
        doc.append("## 🚨 Validation Flags Required Review")
        doc.append("> The AI has detected potential issues that require human review before finalization.\n")
        
        for flag in flags:
            section, f_type, severity, desc = flag
            # High severity usually gets an alert in some MD parsers
            icon = "🔴" if severity == "high" else ("🟡" if severity == "medium" else "🔵")
            doc.append(f"{icon} **[{severity.upper()}] {section.replace('_', ' ').title()} ({f_type})**: {desc}")
        doc.append("\n---\n")
    
    # 2. Compile Sections
    for db_key, display_title in section_order:
        content = sections.get(db_key, "*(Section not generated)*")
        doc.append(f"## {display_title}\n")
        doc.append(content)
        doc.append("\n---\n")
        
    return "\n".join(doc)


def export_brd_to_pdf(session_id: str, output_file: str = None, title: str = "Business Requirements Document") -> bytes:
    """
    Fetches the latest BRD sections and exports as PDF with enhanced formatting.
    
    Automatically detects and styles:
    - Headings (H1-H6) with progressive sizing and colors
    - Bold text (**text** or __text__) in dark color
    - Italic text (*text* or _text_) in gray
    - Code blocks with background color
    - Tables with alternating row colors
    - Blockquotes with left border and color
    - Links with blue color
    
    Args:
        session_id: Session identifier
        output_file: Optional file path to save PDF. If not provided, returns bytes.
        title: Document title
        
    Returns:
        PDF bytes if output_file is None, otherwise None and saves to file
    """
    if not WEASYPRINT_AVAILABLE:
        raise ImportError(
            "weasyprint is required for PDF export. "
            "Install it with: pip install weasyprint"
        )
    
    # Get markdown content
    markdown_content = export_brd(session_id, title)
    
    # Convert markdown to HTML with enhanced markdown extensions
    html_content = markdown.markdown(
        markdown_content,
        extensions=[
            'tables',
            'fenced_code',
            'toc',
            'extra',           # Includes abbreviations, footnotes, definition lists
            'codehilite',      # Better code block styling
        ]
    )
    
    # Post-process to add color support for highlighted keywords
    html_content = _add_color_highlights(html_content)
    
    # Create styled HTML document with comprehensive CSS for all markdown elements
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            * {{
                margin: 0;
                padding: 0;
            }}
            
            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.7;
                color: #2c3e50;
                background-color: #ffffff;
                padding: 40px;
            }}
            
            /* Heading Styles */
            h1 {{
                font-size: 32px;
                color: #0056b3;
                border-bottom: 4px solid #0056b3;
                padding-bottom: 15px;
                margin: 40px 0 30px 0;
                font-weight: 700;
                page-break-after: avoid;
            }}
            
            h2 {{
                font-size: 24px;
                color: #003d82;
                border-left: 6px solid #0056b3;
                padding-left: 15px;
                margin: 30px 0 20px 0;
                font-weight: 700;
                page-break-after: avoid;
            }}
            
            h3 {{
                font-size: 18px;
                color: #1a5490;
                padding-left: 10px;
                margin: 20px 0 15px 0;
                font-weight: 600;
                page-break-after: avoid;
            }}
            
            h4 {{
                font-size: 16px;
                color: #2c5aa0;
                margin: 15px 0 10px 0;
                font-weight: 600;
                page-break-after: avoid;
            }}
            
            h5, h6 {{
                font-size: 14px;
                color: #444;
                margin: 10px 0 8px 0;
                font-weight: 600;
            }}
            
            /* Bold Text */
            strong, b {{
                color: #1a1a1a;
                font-weight: 700;
            }}
            
            /* Italic Text */
            em, i {{
                color: #555;
                font-style: italic;
            }}
            
            /* Paragraph Styles */
            p {{
                margin: 12px 0;
                text-align: justify;
            }}
            
            /* Link Styles */
            a {{
                color: #0056b3;
                text-decoration: underline;
                font-weight: 500;
            }}
            
            /* Code Styles */
            code {{
                background-color: #f5f5f5;
                color: #d63384;
                padding: 4px 8px;
                border-radius: 3px;
                font-family: 'Courier New', 'Consolas', monospace;
                font-size: 13px;
            }}
            
            /* Code Blocks */
            pre {{
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 5px;
                padding: 15px;
                overflow-x: auto;
                border-left: 4px solid #0056b3;
                margin: 15px 0;
                page-break-inside: avoid;
            }}
            
            pre code {{
                background-color: transparent;
                color: #333;
                padding: 0;
                border-radius: 0;
            }}
            
            /* Table Styles */
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
                page-break-inside: avoid;
            }}
            
            th {{
                background-color: #0056b3;
                color: white;
                padding: 12px;
                text-align: left;
                font-weight: 600;
                border: 1px solid #003d82;
            }}
            
            td {{
                padding: 10px 12px;
                border: 1px solid #ddd;
            }}
            
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            
            tr:hover {{
                background-color: #f0f8ff;
            }}
            
            /* Blockquote Styles */
            blockquote {{
                border-left: 5px solid #0056b3;
                padding-left: 20px;
                margin-left: 0;
                margin-right: 0;
                color: #555;
                font-style: italic;
                background-color: #f0f8ff;
                padding: 12px 15px;
                border-radius: 4px;
                margin: 15px 0;
            }}
            
            blockquote p {{
                margin: 5px 0;
            }}
            
            /* List Styles */
            ul, ol {{
                margin: 15px 0;
                padding-left: 30px;
            }}
            
            li {{
                margin: 8px 0;
                line-height: 1.6;
            }}
            
            /* Horizontal Rule */
            hr {{
                border: none;
                border-top: 2px solid #0056b3;
                margin: 30px 0;
            }}
            
            /* Metadata */
            .metadata {{
                background-color: #f0f8ff;
                padding: 12px 15px;
                border-radius: 4px;
                margin-bottom: 20px;
                font-size: 0.9em;
                color: #555;
            }}
            
            /* Highlighted/Important Text */
            .highlight {{
                background-color: #fff3cd;
                padding: 2px 5px;
                border-radius: 3px;
                color: #856404;
            }}
            
            .highlight-critical {{
                background-color: #f8d7da;
                padding: 2px 5px;
                border-radius: 3px;
                color: #721c24;
                font-weight: 600;
            }}
            
            .highlight-success {{
                background-color: #d4edda;
                padding: 2px 5px;
                border-radius: 3px;
                color: #155724;
            }}
            
            .highlight-info {{
                background-color: #d1ecf1;
                padding: 2px 5px;
                border-radius: 3px;
                color: #0c5460;
            }}
            
            /* Page Break Control */
            @page {{
                size: A4;
                margin: 2cm;
            }}
            
            @media print {{
                h1, h2, h3 {{
                    page-break-after: avoid;
                }}
                pre {{
                    page-break-inside: avoid;
                }}
                table {{
                    page-break-inside: avoid;
                }}
                blockquote {{
                    page-break-inside: avoid;
                }}
            }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    # Generate PDF
    pdf_bytes = HTML(string=styled_html).write_pdf()
    
    # Save to file if path provided
    if output_file:
        with open(output_file, 'wb') as f:
            f.write(pdf_bytes)
        return None
    
    return pdf_bytes


def _add_color_highlights(html_content: str) -> str:
    """
    Post-process HTML to add color highlights for important keywords.
    
    Detects patterns like:
    - [CRITICAL: ...]  → Red highlight
    - [SUCCESS: ...]   → Green highlight
    - [INFO: ...]      → Blue highlight
    - [WARNING: ...]   → Yellow highlight (default)
    
    Args:
        html_content: HTML string to process
        
    Returns:
        HTML string with color highlights applied
    """
    # Define highlight patterns and their CSS classes
    patterns = [
        (r'\[CRITICAL:\s*([^\]]+)\]', '<span class="highlight-critical">[CRITICAL: \\1]</span>'),
        (r'\[SUCCESS:\s*([^\]]+)\]', '<span class="highlight-success">[SUCCESS: \\1]</span>'),
        (r'\[INFO:\s*([^\]]+)\]', '<span class="highlight-info">[INFO: \\1]</span>'),
        (r'\[WARNING:\s*([^\]]+)\]', '<span class="highlight">[WARNING: \\1]</span>'),
        (r'\[NOTE:\s*([^\]]+)\]', '<span class="highlight">[NOTE: \\1]</span>'),
    ]
    
    for pattern, replacement in patterns:
        html_content = re.sub(pattern, replacement, html_content, flags=re.IGNORECASE)
    
    # Also detect emoji/icon patterns for icons like 🔴, 🟡, 🔵 and wrap them with spans
    # These usually appear in validation flags
    html_content = re.sub(
        r'(🔴|🟡|🔵|☑️|✓|✗)',
        r'<span style="font-size: 1.2em; margin-right: 5px;">\1</span>',
        html_content
    )
    
    return html_content


def export_brd_to_docx(session_id: str, output_file: str = None, title: str = "Business Requirements Document", template_path: str = None) -> bytes:
    """
    Export BRD as DOCX using template if available, otherwise generate from scratch.
    
    Template-based approach:
    - If template_path provided, uses that template as a base
    - If template_path is None, looks for 'brd.docx' in current directory
    - Falls back to creating DOCX from scratch if no template found
    
    The template should contain placeholder text like:
    - {TITLE} → Document title
    - {SESSION_ID} → Session identifier
    - {GENERATED_DATE} → Generation timestamp
    - {EXECUTIVE_SUMMARY} → Executive summary content
    - {FUNCTIONAL_REQUIREMENTS} → Functional requirements
    - {STAKEHOLDER_ANALYSIS} → Stakeholder analysis
    - {TIMELINE} → Timeline
    - {DECISIONS} → Key decisions
    - {ASSUMPTIONS} → Assertions & assumptions
    - {SUCCESS_METRICS} → Success metrics
    
    Args:
        session_id: Session identifier
        output_file: Optional file path to save DOCX. If not provided, returns bytes.
        title: Document title
        template_path: Optional path to template DOCX file
        
    Returns:
        DOCX bytes if output_file is None, otherwise None and saves to file
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install it with: pip install python-docx"
        )
    
    # Get markdown content
    markdown_content = export_brd(session_id, title)
    
    # Extract sections for template filling
    sections = get_latest_brd_sections(session_id)
    
    # If no template path provided, try to find brd.docx in current directory
    if template_path is None:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(current_dir, "brd.docx")
    
    # Try to use template if it exists
    if os.path.exists(template_path):
        doc = _fill_docx_template(
            template_path, 
            session_id, 
            title, 
            sections
        )
    else:
        # Fallback: Create DOCX from scratch
        doc = _create_docx_from_scratch(
            session_id,
            title,
            sections
        )
    
    # Convert to bytes
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    docx_bytes_content = docx_bytes.read()
    
    # Save to file if path provided
    if output_file:
        with open(output_file, 'wb') as f:
            f.write(docx_bytes_content)
        return None
    
    return docx_bytes_content


def _fill_docx_template(template_path: str, session_id: str, title: str, sections: dict) -> Document:
    """
    Fill a DOCX template with BRD content.
    
    Replaces placeholders in the template with actual content.
    
    Args:
        template_path: Path to the template DOCX file
        session_id: Session identifier
        title: Document title
        sections: Dictionary of BRD sections
        
    Returns:
        Filled Document object
    """
    doc = Document(template_path)
    
    # Prepare replacement dictionary
    replacements = {
        '{TITLE}': title,
        '{SESSION_ID}': session_id,
        '{GENERATED_DATE}': datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC'),
        '{EXECUTIVE_SUMMARY}': sections.get('executive_summary', '(Not generated)'),
        '{FUNCTIONAL_REQUIREMENTS}': sections.get('functional_requirements', '(Not generated)'),
        '{STAKEHOLDER_ANALYSIS}': sections.get('stakeholder_analysis', '(Not generated)'),
        '{TIMELINE}': sections.get('timeline', '(Not generated)'),
        '{DECISIONS}': sections.get('decisions', '(Not generated)'),
        '{ASSUMPTIONS}': sections.get('assumptions', '(Not generated)'),
        '{SUCCESS_METRICS}': sections.get('success_metrics', '(Not generated)'),
    }
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
    
    return doc


def _create_docx_from_scratch(session_id: str, title: str, sections: dict) -> Document:
    """
    Create a DOCX document from scratch (when template not available).
    
    Args:
        session_id: Session identifier
        title: Document title
        sections: Dictionary of BRD sections
        
    Returns:
        Document object
    """
    doc = Document()
    
    # Title
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"Generated: ").bold = True
    metadata.add_run(datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC'))
    
    metadata = doc.add_paragraph()
    metadata.add_run(f"Session ID: ").bold = True
    metadata.add_run(session_id)
    
    doc.add_paragraph()  # Spacing
    
    # Sections
    section_order = [
        ("executive_summary", "1. Executive Summary"),
        ("functional_requirements", "2. Functional Requirements"),
        ("stakeholder_analysis", "3. Stakeholder Analysis"),
        ("timeline", "4. Project Timeline"),
        ("decisions", "5. Key Decisions"),
        ("assumptions", "6. Assertions & Assumptions"),
        ("success_metrics", "7. Success Metrics")
    ]
    
    for db_key, display_title in section_order:
        doc.add_heading(display_title, level=2)
        content = sections.get(db_key, "*(Section not generated)*")
        
        # Add content with basic formatting
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line, style='List Bullet' if line.startswith('-') else 'Normal')
    
    return doc
